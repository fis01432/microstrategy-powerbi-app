import streamlit as st
import pandas as pd
import openai
from openai import OpenAI
import base64
import os
import zipfile
import json
import csv
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from dotenv import load_dotenv

# Configuração inicial
st.set_page_config(layout="wide", page_title="DashMigrate Pro+")
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Diretórios
DATA_DIR = "data"
OUTPUT_DIR = "output"
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Caminhos
CAMINHO_IMAGEM = os.path.join(DATA_DIR, "uploaded_image.png")
CAMINHO_DATASET = os.path.join(DATA_DIR, "user_dataset.xlsx")
CAMINHO_PROGRESSO = os.path.join(DATA_DIR, "progresso.json")
CAMINHO_ROTEIRO = os.path.join(DATA_DIR, "roteiro.json")
CAMINHO_CHECKLIST = os.path.join(DATA_DIR, "checklist.json")
CAMINHO_OCR = os.path.join(DATA_DIR, "ocr_result.json")
CAMINHO_PLATAFORMA = os.path.join(DATA_DIR, "plataforma.json")
CAMINHO_ETAPA_ATUAL = os.path.join(DATA_DIR, "etapa_atual.json")

# Etapas
etapas = [
    "Seleção da plataforma",
    "Upload do dashboard",
    "Extração visual",
    "Validação dos dados",
    "Geração do roteiro",
    "Checklist visual",
    "Exportação final"
]

# Funções auxiliares
def salvar_json(caminho, objeto):
    with open(caminho, "w", encoding="utf-8") as f:
        json.dump(objeto, f, indent=2, ensure_ascii=False)

def carregar_json(caminho):
    if os.path.exists(caminho):
        with open(caminho, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def salvar_etapa_atual(indice):
    salvar_json(CAMINHO_ETAPA_ATUAL, {"indice": indice})

def carregar_etapa_atual():
    return carregar_json(CAMINHO_ETAPA_ATUAL).get("indice", 0)

def resetar_progresso():
    salvar_json(CAMINHO_PROGRESSO, {etapa: False for etapa in etapas})
    salvar_json(CAMINHO_CHECKLIST, {})
    salvar_json(CAMINHO_ROTEIRO, {"conteudo": ""})
    salvar_json(CAMINHO_OCR, {"ocr": ""})
    salvar_json(CAMINHO_ETAPA_ATUAL, {"indice": 0})
    salvar_json(CAMINHO_PLATAFORMA, {"origem": ""})

if not os.path.exists(CAMINHO_PROGRESSO):
    resetar_progresso()

progresso = carregar_json(CAMINHO_PROGRESSO)
checklist_por_componente = carregar_json(CAMINHO_CHECKLIST)
roteiro = carregar_json(CAMINHO_ROTEIRO).get("conteudo", "")
etapa_atual = carregar_etapa_atual()
plataforma_selecionada = carregar_json(CAMINHO_PLATAFORMA).get("origem", "")

# Layout
st.title("📊 Migração Assistida")

# Barra lateral interativa
st.sidebar.image("logo.png", use_container_width=True)
st.sidebar.header("Progresso")

for i, etapa in enumerate(etapas):
    status = progresso.get(etapa, False)
    icone = "✅" if status else "⬜"
    if st.sidebar.button(f"{icone} {etapa}", key=f"etapa_sidebar_{i}"):
        salvar_etapa_atual(i)
        st.rerun()

if plataforma_selecionada:
    st.markdown(f"🧭 Plataforma de origem: **{plataforma_selecionada}**")

# st.header(f"Etapa {etapa_atual + 1}: {etapas[etapa_atual]}")

# Etapa 1: Configuração inicial com layout moderno
if etapa_atual == 0:
    st.markdown("""
    <div style='background-color:#f0f2f6; padding: 20px 30px; border-radius: 12px; margin-bottom: 25px;'>
        <h2 style='color:#1f77b4; margin-bottom: 10px;'>📌 Etapa 1: Configuração Inicial</h2>
        <p style='color:#333;'>Escolha a plataforma de origem, sua versão e o modelo de IA que será utilizado para auxiliar na análise e migração.</p>
    </div>
    """, unsafe_allow_html=True)

    st.progress((etapa_atual + 1) / len(etapas))

    with st.container():
        col1, col2 = st.columns(2)

        with col1:
            plataforma = st.selectbox("🧩 Plataforma de origem:", ["MicroStrategy", "Tableau", "QlikView", "Excel", "Outro"])

        versoes_por_plataforma = {
            "MicroStrategy": ["2021", "2020", "2019", "2018 ou anterior"],
            "Tableau": ["2023.3", "2022.4", "2021.2", "Versão anterior"],
            "QlikView": ["12.6", "12.5", "12.2", "Versão anterior"],
            "Excel": ["Microsoft 365", "Excel 2019", "Excel 2016", "Excel 2013 ou anterior"],
            "Outro": ["Versão desconhecida", "Sem versão definida"]
        }

        with col2:
            versao_origem = st.selectbox(f"🔢 Versão do {plataforma}:", versoes_por_plataforma.get(plataforma, ["Versão desconhecida"]))

    col3, col4 = st.columns(2)

    with col3:
        versao_powerbi = st.selectbox("🟢 Versão do Power BI (destino):", [
            "Power BI Desktop (jun/2024)", "Power BI Desktop (dez/2023)", "Power BI Pro", "Power BI Service"
        ])

    modelos_llm_exibicao = [
        "OpenAI - GPT-4o",
        "OpenAI - GPT-4",
        "OpenAI - GPT-3.5-turbo",
        "Anthropic - Claude 3 Opus",
        "Anthropic - Claude 3 Sonnet",
        "Google - Gemini 1.5 Pro",
        "Google - Gemini 1.0 Pro",
        "Meta - LLaMA 3 (70B)",
        "Meta - LLaMA 3 (8B)",
        "Mistral - Mixtral 8x7B",
        "Mistral - Mistral 7B Instruct",
        "Cohere - Command R+",
        "AWS Bedrock - Titan",
        "Azure OpenAI - GPT-4",
        "Outro modelo externo"
    ]

    modelos_llm_map = {
        "OpenAI - GPT-4o": "gpt-4o",
        "OpenAI - GPT-4": "gpt-4",
        "OpenAI - GPT-3.5-turbo": "gpt-3.5-turbo",
        "Anthropic - Claude 3 Opus": "claude-3-opus",
        "Anthropic - Claude 3 Sonnet": "claude-3-sonnet",
        "Google - Gemini 1.5 Pro": "gemini-1.5-pro",
        "Google - Gemini 1.0 Pro": "gemini-1.0-pro",
        "Meta - LLaMA 3 (70B)": "llama3-70b",
        "Meta - LLaMA 3 (8B)": "llama3-8b",
        "Mistral - Mixtral 8x7b": "mixtral-8x7b",
        "Mistral - Mistral 7B Instruct": "mistral-7b-instruct",
        "Cohere - Command R+": "command-r-plus",
        "AWS Bedrock - Titan": "bedrock-titan",
        "Azure OpenAI - GPT-4": "gpt-4-azure",
        "Outro modelo externo": "custom"
    }

    with col4:
        modelo_exibicao = st.selectbox("🧠 Modelo de LLM:", modelos_llm_exibicao)
        modelo_llm = modelos_llm_map[modelo_exibicao]

    st.markdown("---")

    if st.button("✅ Confirmar configurações e avançar"):
        salvar_json(CAMINHO_PLATAFORMA, {
            "origem": plataforma,
            "versao_origem": versao_origem,
            "versao_destino": versao_powerbi,
            "modelo_llm_exibicao": modelo_exibicao,
            "modelo_llm": modelo_llm
        })
        salvar_etapa_atual(1)
        progresso[etapas[0]] = True
        salvar_json(CAMINHO_PROGRESSO, progresso)
        st.rerun()


# Etapa 2: Upload do dashboard com layout moderno
elif etapa_atual == 1:
    st.markdown("""
    <div style='background-color:#f0f2f6; padding: 20px 30px; border-radius: 12px; margin-bottom: 25px;'>
        <h2 style='color:#1f77b4; margin-bottom: 10px;'>📥 Etapa 2: Upload do Dashboard</h2>
        <p style='color:#333;'>Envie uma imagem representando o dashboard atual na plataforma de origem para que possamos extrair os componentes visuais via IA.</p>
    </div>
    """, unsafe_allow_html=True)

    st.progress((etapa_atual + 1) / len(etapas))

    col1, col2 = st.columns([2, 1])

    with col1:
        img = st.file_uploader("📷 Envie uma imagem do dashboard (PNG ou JPG):", type=["png", "jpg"])
        if img:
            caminho_img_temp = os.path.join(DATA_DIR, "uploaded_image.png")
            with open(caminho_img_temp, "wb") as f:
                f.write(img.read())
            st.success("Imagem enviada com sucesso!")
            st.image(caminho_img_temp, caption="📊 Dashboard carregado", use_container_width=True)

    with col2:
        st.info("💡 Dica: Use screenshots diretas da tela do MicroStrategy, Tableau, Qlik etc. Isso ajuda a IA a entender a estrutura visual com mais precisão.")

    st.markdown("---")

    if img and st.button("➡️ Avançar para extração visual com IA"):
        salvar_etapa_atual(2)
        progresso[etapas[1]] = True
        salvar_json(CAMINHO_PROGRESSO, progresso)
        st.rerun()


# Etapa 3: Extração visual com layout moderno e checklist
elif etapa_atual == 2:
    st.markdown("""
    <div style='background-color:#f0f2f6; padding: 20px 30px; border-radius: 12px; margin-bottom: 25px;'>
        <h2 style='color:#1f77b4; margin-bottom: 10px;'>🔍 Etapa 3: Extração Visual com IA</h2>
        <p style='color:#333;'>Utilizaremos inteligência artificial para extrair automaticamente os elementos visuais do dashboard enviado.</p>
    </div>
    """, unsafe_allow_html=True)

    st.progress((etapa_atual + 1) / len(etapas))

    st.image(CAMINHO_IMAGEM, caption="📷 Imagem do Dashboard enviado", use_container_width=True)

    if st.button("🤖 Executar extração visual com GPT-4o"):
        with open(CAMINHO_IMAGEM, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("utf-8")

        prompt = ("Você é um especialista em BI. Analise a imagem de um dashboard e extraia os elementos visuais, como gráficos, tabelas, indicadores, KPIs, textos, campos, filtros e menus.")
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}}
            ]}],
            temperature=0.3
        )
        resultado = response.choices[0].message.content
        salvar_json(CAMINHO_OCR, {"ocr": resultado})
        progresso[etapas[2]] = True
        salvar_json(CAMINHO_PROGRESSO, progresso)
        salvar_etapa_atual(3)
        st.rerun()

    texto_ocr = carregar_json(CAMINHO_OCR).get("ocr", "")
    if texto_ocr:
        st.markdown("""
        <div style='margin-top: 30px;'>
            <h4 style='color:#1f77b4;'>📄 Relatório gerado pela IA:</h4>
        </div>
        """, unsafe_allow_html=True)
        st.markdown(texto_ocr)

        st.markdown("""
        <div style='margin-top: 20px; background-color:#ffffff; padding:20px; border:1px solid #e1e1e1; border-radius: 10px;'>
        <h5 style='color:#333;'>✅ Checklist de conferência dos elementos detectados</h5>
        <p>Marque os itens confirmados na imagem. Se algum elemento não foi detectado, adicione manualmente abaixo.</p>
        """, unsafe_allow_html=True)

        componentes_extraidos = [linha for linha in texto_ocr.split("\n") if any(x in linha.lower() for x in ["gráfico", "kpi", "filtro", "campo", "tabela"])]
        checklist = st.session_state.get("checklist_etapa3", {})

        for i, comp in enumerate(componentes_extraidos):
            key = f"comp_{i}"
            if key not in checklist:
                checklist[key] = {"texto": comp.strip(), "checado": False}
            checklist[key]["checado"] = st.checkbox(checklist[key]["texto"], checklist[key]["checado"], key=key)

        novo_item = st.text_input("➕ Adicionar elemento manualmente:", key="novo_item_input")
        if st.button("Adicionar item", key="adicionar_item_btn") and novo_item.strip():
            novo_key = f"manual_{len(checklist)}"
            checklist[novo_key] = {"texto": novo_item.strip(), "checado": False}
            st.rerun()

        st.session_state["checklist_etapa3"] = checklist

        if st.button("➡️ Avançar para validação de dados"):
            progresso[etapas[2]] = True
            salvar_json(CAMINHO_PROGRESSO, progresso)
            salvar_etapa_atual(3)
            st.rerun()


# Etapa 4: Validação de dados com layout moderno e conexão flexível
elif etapa_atual == 3:
    st.markdown("""
    <div style='background-color:#f9f9f9; padding: 25px; border-radius: 12px;'>
        <h2 style='color:#1f77b4;'>🔗 Etapa 4: Validação de Dados</h2>
        <p>Escolha como deseja fornecer os dados para que possamos avaliar a compatibilidade com o dashboard extraído.</p>
    </div>
    """, unsafe_allow_html=True)

    modo_dados = st.radio("Fonte de dados:", ["📁 Upload de arquivo", "🔌 Conectar a banco ou nuvem"])

    df = None
    colunas_disponiveis = []

    if modo_dados == "📁 Upload de arquivo":
        file = st.file_uploader("📁 Envie a base de dados (.xlsx ou .csv)", type=["xlsx", "csv"])
        if file:
            caminho_base = os.path.join(DATA_DIR, file.name)
            with open(caminho_base, "wb") as f:
                f.write(file.read())

            try:
                if file.name.endswith(".csv"):
                    df = pd.read_csv(caminho_base)
                else:
                    df = pd.read_excel(caminho_base, engine="openpyxl")
                colunas_disponiveis = list(df.columns)
                st.success("✅ Base carregada com sucesso!")
                st.markdown("### 📊 Pré-visualização dos dados")
                st.dataframe(df.head())
            except Exception as e:
                st.error(f"Erro ao carregar base: {e}")

    elif modo_dados == "🔌 Conectar a banco ou nuvem":
        tipo_conexao = st.selectbox("Tipo de fonte de dados", [
            "MySQL", "PostgreSQL", "SQL Server",
            "Databricks (JDBC URL)",
            "AWS S3 (CSV ou Parquet)",
            "Azure Blob (CSV ou Parquet)"
        ])

        # Conexões tradicionais
        if tipo_conexao in ["MySQL", "PostgreSQL", "SQL Server"]:
            host = st.text_input("Host")
            porta = st.text_input("Porta", value="3306" if tipo_conexao == "MySQL" else "5432")
            database = st.text_input("Nome do banco de dados")
            usuario = st.text_input("Usuário")
            senha = st.text_input("Senha", type="password")
            tabela = st.text_input("Nome da tabela")

            if st.button("Conectar e carregar dados"):
                try:
                    if tipo_conexao == "MySQL":
                        import pymysql
                        conn = pymysql.connect(host=host, user=usuario, password=senha, database=database, port=int(porta))
                    elif tipo_conexao == "PostgreSQL":
                        import psycopg2
                        conn = psycopg2.connect(host=host, user=usuario, password=senha, dbname=database, port=int(porta))
                    elif tipo_conexao == "SQL Server":
                        import pyodbc
                        conn_str = f"DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={host},{porta};DATABASE={database};UID={usuario};PWD={senha}"
                        conn = pyodbc.connect(conn_str)

                    df = pd.read_sql(f"SELECT * FROM {tabela}", conn)
                    conn.close()
                    colunas_disponiveis = list(df.columns)
                    st.success("✅ Dados carregados com sucesso!")
                    st.dataframe(df.head())
                except Exception as e:
                    st.error(f"Erro ao conectar: {e}")

        elif tipo_conexao == "Databricks (JDBC URL)":
            jdbc_url = st.text_input("JDBC URL")
            token = st.text_input("Token pessoal do Databricks", type="password")
            query = st.text_area("Consulta SQL", "SELECT * FROM nome_tabela")

            if st.button("Conectar ao Databricks"):
                try:
                    import sqlalchemy
                    engine = sqlalchemy.create_engine(f"{jdbc_url};AuthMech=3;UID=token;PWD={token}")
                    df = pd.read_sql(query, con=engine)
                    colunas_disponiveis = list(df.columns)
                    st.success("✅ Dados carregados do Databricks!")
                    st.dataframe(df.head())
                except Exception as e:
                    st.error(f"Erro ao conectar: {e}")

        elif tipo_conexao == "AWS S3 (CSV ou Parquet)":
            access_key = st.text_input("AWS Access Key")
            secret_key = st.text_input("AWS Secret Key", type="password")
            bucket = st.text_input("Bucket")
            caminho_arquivo = st.text_input("Caminho do arquivo (ex: pasta/arquivo.csv)")
            tipo_arquivo = st.selectbox("Tipo de arquivo", ["csv", "parquet"])

            if st.button("Conectar ao S3"):
                try:
                    import boto3, io
                    s3 = boto3.client('s3', aws_access_key_id=access_key, aws_secret_access_key=secret_key)
                    obj = s3.get_object(Bucket=bucket, Key=caminho_arquivo)
                    if tipo_arquivo == "csv":
                        df = pd.read_csv(obj['Body'])
                    else:
                        import pyarrow.parquet as pq
                        df = pq.read_table(obj['Body']).to_pandas()
                    colunas_disponiveis = list(df.columns)
                    st.success("✅ Arquivo carregado do S3!")
                    st.dataframe(df.head())
                except Exception as e:
                    st.error(f"Erro ao acessar S3: {e}")

        elif tipo_conexao == "Azure Blob (CSV ou Parquet)":
            conn_str = st.text_input("Connection String do Azure")
            container = st.text_input("Nome do container")
            blob = st.text_input("Caminho do arquivo")
            tipo_arquivo = st.selectbox("Tipo de arquivo", ["csv", "parquet"])

            if st.button("Conectar ao Azure Blob"):
                try:
                    from azure.storage.blob import BlobServiceClient
                    import io
                    blob_service_client = BlobServiceClient.from_connection_string(conn_str)
                    blob_client = blob_service_client.get_blob_client(container=container, blob=blob)
                    stream = blob_client.download_blob().readall()
                    if tipo_arquivo == "csv":
                        df = pd.read_csv(io.BytesIO(stream))
                    else:
                        import pyarrow.parquet as pq
                        df = pq.read_table(io.BytesIO(stream)).to_pandas()
                    colunas_disponiveis = list(df.columns)
                    st.success("✅ Arquivo carregado do Azure Blob!")
                    st.dataframe(df.head())
                except Exception as e:
                    st.error(f"Erro ao acessar Azure Blob: {e}")

    # Análise com IA
    if df is not None and len(colunas_disponiveis) > 0:
        texto_ocr = carregar_json(CAMINHO_OCR).get("ocr", "")

        st.markdown("<br><h4>🤖 Análise de Compatibilidade com o Dashboard</h4>", unsafe_allow_html=True)
        with st.spinner("Analisando com inteligência artificial..."):
            prompt = f"""
Você é um consultor de BI. Um dashboard foi extraído visualmente com o seguinte conteúdo:

{texto_ocr}

A seguir, temos uma base de dados com estas colunas:
{', '.join(colunas_disponiveis)}

Sua tarefa:
1. Verifique quais gráficos, KPIs, filtros e tabelas descritos no dashboard podem ser construídos com os dados disponíveis.
2. Apresente uma tabela listando o nome do componente, os campos correspondentes e se está compatível (Sim/Não).
3. Quando algum campo não for encontrado, sugira como ele pode ser criado (ex: "Ticket Médio = receita / quantidade").
4. Apresente tudo de forma clara para o analista entender rapidamente o que pode ser implementado e o que falta.
"""
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2
            )
            analise_dados = response.choices[0].message.content
            st.markdown(analise_dados)

        if st.button("➡️ Avançar para geração do roteiro técnico"):
            progresso[etapas[3]] = True
            salvar_json(CAMINHO_PROGRESSO, progresso)
            salvar_etapa_atual(4)
            st.rerun()




# Etapa 5: Geração do roteiro técnico detalhado com layout moderno e medidas DAX
elif etapa_atual == 4:
    st.markdown("""
    <div style='background-color:#f5f5f5; padding: 25px; border-radius: 12px;'>
        <h2 style='color:#1f77b4;'>🛠️ Etapa 5: Geração do Roteiro Técnico</h2>
        <p>Com base na imagem do dashboard extraído e nos dados carregados, vamos gerar um roteiro técnico detalhado, incluindo as medidas DAX necessárias para o Power BI.</p>
    </div>
    """, unsafe_allow_html=True)

    texto_ocr = carregar_json(CAMINHO_OCR).get("ocr", "")
    colunas_disponiveis = []

    try:
        arquivos = os.listdir(DATA_DIR)
        planilhas = [f for f in arquivos if f.endswith((".csv", ".xlsx"))]
        if planilhas:
            caminho_arquivo = os.path.join(DATA_DIR, planilhas[0])
            if caminho_arquivo.endswith(".csv"):
                df_base = pd.read_csv(caminho_arquivo)
            else:
                df_base = pd.read_excel(caminho_arquivo, engine="openpyxl")
            colunas_disponiveis = list(df_base.columns)
    except Exception as e:
        st.warning("⚠️ Não foi possível ler a base de dados.")

    if st.button("🚀 Gerar roteiro técnico completo"):
        with st.spinner("Gerando roteiro com instruções e medidas DAX..."):
            prompt_roteiro = f"""
Você é um especialista em BI migrando dashboards do MicroStrategy para o Power BI.

Dashboard extraído visualmente:
{texto_ocr}

Base de dados com colunas:
{', '.join(colunas_disponiveis)}

Gere um roteiro completo com:
1. Conexão da base
2. Transformações necessárias (Power Query)
3. Layout visual do dashboard
4. Tipo de gráficos, eixos, valores e filtros
5. Interações como drill-down e slicers
6. Onde e como posicionar os elementos visuais
7. Recomendações práticas para performance e usabilidade
"""
            response_roteiro = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt_roteiro}],
                temperature=0.2
            )
            roteiro = response_roteiro.choices[0].message.content

            prompt_dax = f"""
Você é um especialista em Power BI.

Com base no seguinte dashboard (extraído do MicroStrategy):
{texto_ocr}

E nesta base de dados com colunas:
{', '.join(colunas_disponiveis)}

Gere as principais **medidas DAX** que devem ser criadas no Power BI.

Para cada medida informe:
- Nome da Medida
- Fórmula DAX
- Descrição do que ela faz
- Onde ela deve ser usada (gráfico, KPI, filtro etc.)
"""
            response_dax = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt_dax}],
                temperature=0.2
            )
            medidas_dax = response_dax.choices[0].message.content

            roteiro_completo = f"""
## 📘 Roteiro Técnico

{roteiro}

---

## 🧮 Medidas DAX Recomendadas

{medidas_dax}
"""
            salvar_json(CAMINHO_ROTEIRO, {"conteudo": roteiro_completo})
            progresso[etapas[4]] = True
            salvar_json(CAMINHO_PROGRESSO, progresso)
            salvar_etapa_atual(5)
            st.rerun()

    roteiro_salvo = carregar_json(CAMINHO_ROTEIRO).get("conteudo", "")
    if roteiro_salvo:
        st.subheader("📄 Roteiro Técnico + Medidas DAX")
        st.markdown(roteiro_salvo, unsafe_allow_html=True)

        if st.button("➡️ Avançar para checklist visual"):
            salvar_etapa_atual(5)
            st.rerun()



# Etapa 6: Checklist visual com comparação de imagens e elementos modernizado
elif etapa_atual == 5:
    st.markdown("""
    <div style='background-color:#f5f5f5; padding: 25px; border-radius: 12px;'>
        <h2 style='color:#1f77b4;'>🖼️ Etapa 6: Checklist Visual e Comparação de Dashboards</h2>
        <p>Nesta etapa, vamos comparar visualmente o dashboard original do MicroStrategy com o novo criado no Power BI. A IA auxiliará na verificação de compatibilidade dos elementos.</p>
    </div>
    """, unsafe_allow_html=True)

    texto_ocr = carregar_json(CAMINHO_OCR).get("ocr", "")
    roteiro = carregar_json(CAMINHO_ROTEIRO).get("conteudo", "")
    checklist_por_componente = carregar_json(CAMINHO_CHECKLIST)
    caminho_img_original = CAMINHO_IMAGEM

    uploaded_powerbi_img = st.file_uploader("📤 Envie o screenshot do dashboard recriado no Power BI", type=["png", "jpg"], key="powerbi_img")

    if uploaded_powerbi_img:
        caminho_img_nova = os.path.join(DATA_DIR, "powerbi_dashboard.png")
        with open(caminho_img_nova, "wb") as f:
            f.write(uploaded_powerbi_img.read())

        st.markdown("### 🔍 Visualização lado a lado")
        col1, col2 = st.columns(2)
        with col1:
            st.image(caminho_img_original, caption="Original - MicroStrategy", use_container_width=True)
        with col2:
            st.image(caminho_img_nova, caption="Novo - Power BI", use_container_width=True)

        if st.button("🔎 Comparar dashboards e gerar checklist"):
            with open(caminho_img_original, "rb") as f1, open(caminho_img_nova, "rb") as f2:
                img_original_b64 = base64.b64encode(f1.read()).decode("utf-8")
                img_nova_b64 = base64.b64encode(f2.read()).decode("utf-8")

            prompt = f"""
Você é um consultor de BI. Compare as imagens de dois dashboards:

- O primeiro foi gerado no MicroStrategy (imagem 1)
- O segundo foi recriado no Power BI (imagem 2)

Objetivo:
1. Verifique se os elementos visuais (gráficos, KPIs, filtros, layout) foram mantidos.
2. Liste os componentes como "Compatível", "Parcial" ou "Incompatível".
3. Quando houver discrepâncias, explique o motivo e gere instruções para o analista corrigir no Power BI.
4. Considere também o roteiro abaixo, que foi usado como base para construir o novo dashboard:

{roteiro}

Responda com uma tabela comparativa seguida das instruções de correção.
"""

            with st.spinner("Executando comparação visual com IA..."):
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role": "user", "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_original_b64}"}},
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_nova_b64}"}}
                    ]}],
                    temperature=0.3
                )
                analise_final = response.choices[0].message.content
                checklist_por_componente["comparacao_visual_final"] = analise_final
                salvar_json(CAMINHO_CHECKLIST, checklist_por_componente)
                st.rerun()

    if "comparacao_visual_final" in checklist_por_componente:
        st.subheader("📋 Resultado da Análise Visual")
        st.markdown(checklist_por_componente["comparacao_visual_final"])

        if st.button("📤 Avançar para exportação final"):
            progresso[etapas[5]] = True
            salvar_json(CAMINHO_PROGRESSO, progresso)
            salvar_etapa_atual(6)
            st.rerun()



# Etapa 7: Exportação Final com Relatório Executivo Modernizado
elif etapa_atual == 6:
    st.markdown("""
    <div style='background-color:#e8f4ff; padding: 25px; border-radius: 12px;'>
        <h2 style='color:#1f77b4;'>📤 Etapa Final: Exportação e Relatório Executivo</h2>
        <p>Parabéns! Você chegou à etapa final da migração.</p>
        <p>Agora você pode gerar um relatório executivo contendo:</p>
        <ul>
            <li>Status de todas as etapas</li>
            <li>Comparação visual dos dashboards</li>
            <li>Roteiro técnico completo</li>
            <li>Checklist de verificação e instruções de correção</li>
        </ul>
    </div>
    """, unsafe_allow_html=True)

    if st.button("📄 Gerar Relatório Executivo (.docx)"):
        from docx import Document
        from docx.shared import Inches
        from datetime import datetime

        caminho_img_original = os.path.join(DATA_DIR, "uploaded_image.png")
        caminho_img_novo = os.path.join(DATA_DIR, "powerbi_dashboard.png")
        ocr = carregar_json(CAMINHO_OCR).get("ocr", "")
        roteiro = carregar_json(CAMINHO_ROTEIRO).get("conteudo", "")
        checklist = carregar_json(CAMINHO_CHECKLIST)
        progresso = carregar_json(CAMINHO_PROGRESSO)

        doc = Document()
        doc.add_heading("Relatório Executivo de Migração de Dashboard", 0)
        doc.add_paragraph(f"Projeto: DashMigrate Pro+")
        doc.add_paragraph(f"Data da Migração: {datetime.now().strftime('%d/%m/%Y')}")
        doc.add_paragraph("Plataforma Origem: MicroStrategy")
        doc.add_paragraph("Plataforma Destino: Power BI")
        doc.add_paragraph("Analista Responsável: __________________________")
        doc.add_page_break()

        doc.add_heading("Resumo Executivo", level=1)
        doc.add_paragraph("Este relatório apresenta a análise completa da migração de um dashboard da plataforma MicroStrategy para o Power BI, incluindo os resultados da extração visual, validação dos dados, roteiro técnico, comparação visual e checklist final de qualidade.")

        doc.add_heading("Status das Etapas", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Etapa'
        hdr_cells[1].text = 'Status'
        for etapa, status in progresso.items():
            row = table.add_row().cells
            row[0].text = etapa
            row[1].text = "✅ Completa" if status else "⏳ Pendente"
        doc.add_page_break()

        doc.add_heading("Comparação Visual dos Dashboards", level=1)
        doc.add_paragraph("As imagens abaixo representam o dashboard original (MicroStrategy) e o novo dashboard criado no Power BI.")
        if os.path.exists(caminho_img_original):
            doc.add_picture(caminho_img_original, width=Inches(5.5))
            doc.add_paragraph("🔼 Dashboard Original (MicroStrategy)")
        if os.path.exists(caminho_img_novo):
            doc.add_picture(caminho_img_novo, width=Inches(5.5))
            doc.add_paragraph("🔼 Dashboard Novo (Power BI)")
        doc.add_page_break()

        doc.add_heading("Componentes Migrados - Visão Inicial", level=1)
        doc.add_paragraph("Resumo dos componentes identificados no dashboard original via OCR:")
        doc.add_paragraph(ocr)
        doc.add_page_break()

        doc.add_heading("Roteiro Técnico de Migração", level=1)
        doc.add_paragraph(roteiro)
        doc.add_page_break()

        doc.add_heading("Checklist de Verificação e Comparação Final", level=1)
        comparacao = checklist.get("comparacao_visual_final", checklist.get("analise_comparativa", ""))
        doc.add_paragraph(comparacao)
        doc.add_page_break()

        doc.add_heading("Considerações Finais", level=1)
        doc.add_paragraph(
            "A migração apresenta alto grau de compatibilidade. Recomenda-se revisar campos marcados como parcialmente compatíveis ou incompatíveis e aplicar os ajustes listados no roteiro técnico. "
            "A abordagem automatizada do DashMigrate Pro+ facilita a reprodutibilidade, reduz erros e padroniza o processo de migração de dashboards."
        )

        caminho_saida = os.path.join("output", "relatorio_executivo_dashmigrate.docx")
        os.makedirs("output", exist_ok=True)
        doc.save(caminho_saida)

        with open(caminho_saida, "rb") as file:
            st.download_button(
                label="📥 Baixar Relatório Executivo (.docx)",
                data=file,
                file_name="relatorio_executivo_dashmigrate.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    if st.button("✅ Encerrar e avaliar experiência"):
        salvar_etapa_atual(7)
        st.rerun()

# Etapa 8: Avaliação com layout visual e estrelas clicáveis
elif etapa_atual == 7:
    st.markdown("""
    <div style='background-color:#f0f8ff; padding: 25px; border-radius: 12px;'>
        <h2 style='color:#1f77b4;'>📝 Avaliação da Experiência </h2>
        <p>Agora que você concluiu todo o processo de migração, gostaríamos de saber como foi sua experiência com a plataforma.</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns([1, 3])
    with col1:
        st.markdown("<b>Você recomendaria a plataforma?</b>", unsafe_allow_html=True)
    with col2:
        recomendaria = st.radio("", ["👍 Sim", "👎 Não"], horizontal=True)

    st.markdown("""<b>Avaliação geral (1-Ruim/5-Excelente)</b>""", unsafe_allow_html=True)
    estrelas = st.slider("Qual nota você dá para a experiência geral?", 1, 5, 4)

    st.markdown("### 💬 Comentários ou sugestões")
    comentario = st.text_area("Deixe aqui sua opinião sincera ou sugestões de melhoria")

    if st.button("📩 Enviar avaliação"):
        from datetime import datetime
        import csv

        dados_avaliacao = {
            "data": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "recomendaria": recomendaria,
            "nota_geral": estrelas,
            "comentario": comentario,
        }

        caminho_avaliacao = os.path.join("output", "avaliacoes.csv")
        os.makedirs("output", exist_ok=True)
        escrever_cabecalho = not os.path.exists(caminho_avaliacao)
        with open(caminho_avaliacao, mode="a", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=dados_avaliacao.keys())
            if escrever_cabecalho:
                writer.writeheader()
            writer.writerow(dados_avaliacao)

        st.success("✅ Obrigado pela sua avaliação! Seus dados foram registrados com sucesso.")

        # if st.button("🔄 Voltar ao início"):
        #     resetar_progresso()
        #     st.rerun()

        if st.button("🔄 Voltar ao início"):
            resetar_progresso()
            salvar_etapa_atual(0)  # <- importante para garantir
            st.session_state.clear()  # <- limpa sessão inteira (inclui etapa_atual em memória)
            st.rerun()


